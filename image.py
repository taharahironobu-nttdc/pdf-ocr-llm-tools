import base64
import json
import os
import sys
from pathlib import Path

import requests

# 利用可能なVisionモデル
VISION_MODELS = {
    # 無料モデル (推奨)
    "llama-vision-free": "meta-llama/llama-3.2-11b-vision-instruct:free",
    "qwen-32b-free": "qwen/qwen2.5-vl-32b-instruct:free",
    "qwen-72b-free": "qwen/qwen2.5-vl-72b-instruct:free",
    "kimi-vl-free": "moonshotai/kimi-vl-a3b-thinking:free",
    "gemma-4b-free": "google/gemma-3-4b-it:free",
    "gemma-12b-free": "google/gemma-3-12b-it:free",
    "gemma-27b-free": "google/gemma-3-27b-it:free",
    # 有料モデル
    "llama-3.2-90b": "meta-llama/llama-3.2-90b-vision-instruct",
    "llama-3.2-11b": "meta-llama/llama-3.2-11b-vision-instruct",
    "gpt-4-vision": "openai/gpt-4-vision-preview",
    "gpt-4o": "openai/gpt-4o",
    "claude-3.5-sonnet": "anthropic/claude-3.5-sonnet",
    "gemini-1.5-pro": "google/gemini-pro-1.5",
    "pixtral-12b": "mistralai/pixtral-12b-2409",
}


def convert_pdf_to_images(pdf_path, dpi=200):
    """PDFを画像に変換して、PDFファイル名のディレクトリに保存"""
    try:
        import fitz  # PyMuPDF

        # PDFファイル名からディレクトリ名を生成
        pdf_name = Path(pdf_path).stem
        output_dir = f"{pdf_name}_images"

        os.makedirs(output_dir, exist_ok=True)
        pdf = fitz.open(pdf_path)
        image_paths = []

        print(f"PDFを画像に変換中: {len(pdf)} ページ")
        print(f"保存先: {output_dir}/")

        for page_num, page in enumerate(pdf):
            # 高解像度で変換
            zoom = dpi / 72  # 72 DPI base
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            image_path = f"{output_dir}/page_{page_num + 1:03d}.png"
            pix.save(image_path)
            image_paths.append(image_path)
            print(f"  ページ {page_num + 1}/{len(pdf)} 変換完了 -> {image_path}")

        pdf.close()
        return output_dir, image_paths
    except Exception as e:
        print(f"PDF変換エラー: {e}")
        return None, None


def encode_image(image_path):
    """画像をBase64エンコード"""
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")


def ocr_with_openrouter(image_path, model, api_key):
    """OpenRouterを使ってOCR実行"""
    try:
        # 画像をBase64エンコード
        base64_image = encode_image(image_path)

        # OpenRouter API呼び出し
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "HTTP-Referer": "https://github.com/pdf-converter",
                "X-Title": "PDF to Word Converter",
            },
            json={
                "model": model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """この画像からテキストを正確に抽出してください。

要件:
- レイアウト、表、箇条書きなどの構造を保持
- 出力はマークダウン形式
- 余計な説明は不要、テキストのみを出力
- 日本語の場合は日本語で、英語の場合は英語で出力""",
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}"
                                },
                            },
                        ],
                    }
                ],
                "max_tokens": 4000,
                "temperature": 0.1,
            },
            timeout=120,
        )

        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
        else:
            error_msg = response.json().get("error", {})
            print(f"OpenRouter APIエラー: {response.status_code}")
            print(f"詳細: {error_msg}")
            return None

    except requests.exceptions.RequestException as e:
        print(f"リクエストエラー: {e}")
        return None
    except Exception as e:
        print(f"OCRエラー: {e}")
        return None


def markdown_to_docx(markdown_text, output_path):
    """マークダウンをWordドキュメントに変換"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        doc = Document()

        lines = markdown_text.split("\n")
        in_code_block = False
        in_table = False
        table_data = []

        for line in lines:
            # コードブロックの処理
            if line.strip().startswith("```"):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                p = doc.add_paragraph(line)
                p.style = "Normal"
                run = p.runs[0]
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                continue

            # テーブルの処理
            if "|" in line and line.strip().startswith("|"):
                cells = [cell.strip() for cell in line.split("|")[1:-1]]
                if cells:
                    if not in_table:
                        in_table = True
                        table_data = [cells]
                    elif all(c.replace("-", "").strip() == "" for c in cells):
                        # テーブルヘッダー区切り行をスキップ
                        continue
                    else:
                        table_data.append(cells)
                continue
            elif in_table:
                # テーブル終了、Word tableとして追加
                if len(table_data) > 0:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = "Table Grid"
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            table.rows[i].cells[j].text = cell_data
                table_data = []
                in_table = False

            # 見出しの処理
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            # 箇条書きの処理
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                doc.add_paragraph(line.strip()[2:], style="List Bullet")
            # 番号付きリストの処理
            elif line.strip() and line.strip()[0].isdigit() and ". " in line:
                text = (
                    line.strip().split(". ", 1)[1]
                    if ". " in line.strip()
                    else line.strip()
                )
                doc.add_paragraph(text, style="List Number")
            # 太字の処理
            elif "**" in line:
                p = doc.add_paragraph()
                parts = line.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # 奇数インデックスは太字
                        run.bold = True
            # 通常のテキスト
            elif line.strip():
                doc.add_paragraph(line)
            else:
                # 空行
                doc.add_paragraph()

        # 最後にテーブルが残っている場合
        if in_table and table_data:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = "Table Grid"
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_data

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Word変換エラー: {e}")
        import traceback

        traceback.print_exc()
        return False


def cleanup_temp_files(image_dir, image_paths):
    """一時ファイルを削除"""
    print(f"\n一時ファイルを削除中: {image_dir}/")
    for img_path in image_paths:
        try:
            os.remove(img_path)
        except:
            pass

    try:
        os.rmdir(image_dir)
        print(f"ディレクトリを削除: {image_dir}/")
    except:
        pass


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="PDFをWordに変換 (OpenRouter使用)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
利用可能なモデル:

🆓 無料モデル (おすすめ):
  qwen-72b-free        - Qwen2.5-VL 72B (最も高性能、無料)
  qwen-32b-free        - Qwen2.5-VL 32B (バランス型、無料)
  llama-vision-free    - Llama 3.2 11B Vision (汎用、無料)
  kimi-vl-free         - Kimi VL A3B (軽量、無料)
  gemma-27b-free       - Gemma 3 27B (Google製、無料)
  gemma-12b-free       - Gemma 3 12B (中サイズ、無料)
  gemma-4b-free        - Gemma 3 4B (超軽量、無料)

💰 有料モデル:
  llama-3.2-90b        - Llama 3.2 90B Vision
  gpt-4o               - GPT-4o (OpenAI)
  claude-3.5-sonnet    - Claude 3.5 Sonnet (Anthropic)
  gemini-1.5-pro       - Gemini 1.5 Pro (Google)

おすすめの使い分け:
  - 日本語OCR: qwen-72b-free または qwen-32b-free
  - 英語OCR: llama-vision-free または gemma-27b-free
  - 高速処理: kimi-vl-free または gemma-4b-free
""",
    )
    parser.add_argument("input_pdf", help="入力PDFファイル")
    parser.add_argument("output_docx", nargs="?", help="出力Wordファイル（省略可）")
    parser.add_argument(
        "--model",
        default="qwen-72b-free",
        choices=list(VISION_MODELS.keys()),
        help="使用するモデル (デフォルト: qwen-72b-free, 無料)",
    )
    parser.add_argument("--api-key", help="OpenRouter APIキー")
    parser.add_argument(
        "--keep-images",
        action="store_true",
        help="画像ファイルを保持 (デフォルトで保持されます)",
    )
    parser.add_argument(
        "--delete-images", action="store_true", help="変換後に画像ファイルを削除"
    )
    parser.add_argument(
        "--dpi", type=int, default=200, help="画像変換時のDPI (デフォルト: 200)"
    )
    parser.add_argument(
        "--list-models", action="store_true", help="利用可能なモデルを表示"
    )

    args = parser.parse_args()

    if args.list_models:
        print("\n利用可能なモデル:")
        for short_name, full_name in VISION_MODELS.items():
            print(f"  {short_name:20s} -> {full_name}")
        return

    # APIキーの取得
    api_key = args.api_key or os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        print("エラー: OPENROUTER_API_KEYが設定されていません")
        print("\n以下のいずれかの方法で設定してください:")
        print("1. export OPENROUTER_API_KEY='your-api-key'")
        print("2. python convert_openrouter.py input.pdf --api-key your-api-key")
        sys.exit(1)

    input_pdf = args.input_pdf
    if not os.path.exists(input_pdf):
        print(f"エラー: ファイルが見つかりません: {input_pdf}")
        sys.exit(1)

    # 出力ファイル名を決定
    if args.output_docx:
        output_docx = args.output_docx
    else:
        base_name = Path(input_pdf).stem
        output_docx = f"{base_name}.docx"

    model_full_name = VISION_MODELS[args.model]

    print(f"\n=== PDF to Word 変換 (OpenRouter) ===")
    print(f"入力: {input_pdf}")
    print(f"出力: {output_docx}")
    print(f"モデル: {args.model} ({model_full_name})")
    print(f"DPI: {args.dpi}\n")

    # Step 1: PDFを画像に変換
    print("Step 1: PDFを画像に変換中...")
    image_dir, image_paths = convert_pdf_to_images(input_pdf, dpi=args.dpi)
    if not image_paths:
        print("✗ PDF変換失敗")
        sys.exit(1)

    print(f"\n画像保存先: {image_dir}/")
    print(f"画像ファイル数: {len(image_paths)}\n")

    # Step 2: 各ページをOCR
    print("\nStep 2: OCR実行中...")
    all_text = []

    for i, img_path in enumerate(image_paths, 1):
        print(f"  ページ {i}/{len(image_paths)} 処理中...")
        text = ocr_with_openrouter(img_path, model_full_name, api_key)
        if text:
            all_text.append(text)
            if i < len(image_paths):
                all_text.append("\n\n---\n\n")  # ページ区切り
        else:
            print(f"  ⚠ ページ {i} のOCR失敗")

    if not all_text:
        print("\n✗ 全ページのOCR失敗")
        if args.delete_images:
            cleanup_temp_files(image_dir, image_paths)
        sys.exit(1)

    # Step 3: Wordドキュメントに変換
    print("\nStep 3: Word文書に変換中...")
    combined_text = "".join(all_text)

    if markdown_to_docx(combined_text, output_docx):
        print(f"\n✓ 変換成功: {output_docx}")
    else:
        print("\n⚠ Word変換に失敗しました")
        # マークダウンファイルとして保存
        md_output = output_docx.replace(".docx", ".md")
        with open(md_output, "w", encoding="utf-8") as f:
            f.write(combined_text)
        print(f"  マークダウンとして保存: {md_output}")

    # Step 4: クリーンアップ（オプション）
    if args.delete_images:
        cleanup_temp_files(image_dir, image_paths)
    else:
        print(f"\n📁 画像ファイル保存: {image_dir}/")
        print(f"   ファイル数: {len(image_paths)}")

    print("\n✅ 完了!")


if __name__ == "__main__":
    main()
